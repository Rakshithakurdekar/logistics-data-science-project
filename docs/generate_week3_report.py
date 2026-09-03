from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from generate_week1_report import add_bullet, add_code, add_table


ROOT = Path(__file__).resolve().parents[1]
CHART_DIR = ROOT / "visualisations" / "week3"
OUTPUT = ROOT / "docs" / "week3_advanced_analysis_report.docx"


def build_dataset(seed=42, rows=1000):
    rng = np.random.default_rng(seed)
    carriers = rng.choice(["NorthStar", "SwiftHaul", "MetroFreight", "BlueLine"], rows, p=[0.28, 0.27, 0.25, 0.20])
    regions = rng.choice(["North", "South", "East", "West"], rows)
    service = rng.choice(["Standard", "Express", "Economy"], rows, p=[0.48, 0.27, 0.25])
    dates = pd.date_range("2025-01-01", periods=180, freq="D")
    order_date = pd.Series(rng.choice(dates, rows))
    distance = np.clip(rng.gamma(3.0, 55.0, rows), 8, 520)
    weight = np.clip(rng.lognormal(2.6, 0.65, rows), 1, 70)
    congestion = rng.choice([0, 1], rows, p=[0.76, 0.24])
    carrier_effect = pd.Series(carriers).map({"NorthStar": 0.2, "SwiftHaul": 0.7, "MetroFreight": 0.4, "BlueLine": 1.1}).to_numpy()
    service_effect = pd.Series(service).map({"Standard": 1.6, "Express": 0.7, "Economy": 2.6}).to_numpy()
    delay = np.maximum(0, rng.normal(0.35 + distance / 260 + weight / 55 + congestion * 1.2 + carrier_effect, 0.8, rows))
    delivery_time = np.maximum(1, service_effect + distance / 180 + delay + rng.normal(0, 0.45, rows))
    cost = np.maximum(12, 18 + distance * 0.34 + weight * 1.8 + delivery_time * 5.5 + rng.normal(0, 9, rows))
    status = np.where(delay > 1.25, "Delayed", "On Time")
    data = pd.DataFrame({
        "order_date": order_date, "carrier": carriers, "region": regions, "service_level": service,
        "distance_km": distance.round(1), "weight_kg": weight.round(1), "congestion_flag": congestion,
        "delivery_time_days": delivery_time.round(2), "delay_days": delay.round(2), "shipping_cost": cost.round(2),
        "delivery_status": status,
    })
    return data


def save_charts(data):
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({"axes.grid": True, "grid.alpha": 0.25, "font.size": 9})
    chart_paths = []

    path = CHART_DIR / "delivery_time_distribution.png"
    fig, ax = plt.subplots(figsize=(8, 4.5))
    colors = {"Standard": "#1F4E79", "Express": "#2A9D8F", "Economy": "#E76F51"}
    for service in ["Standard", "Express", "Economy"]:
        ax.hist(data.loc[data["service_level"] == service, "delivery_time_days"], bins=28, histtype="step", linewidth=2, color=colors[service], label=service)
    ax.legend(title="Service level")
    ax.set(title="Delivery-time distribution by service level", xlabel="Delivery time (days)", ylabel="Shipments")
    fig.tight_layout(); fig.savefig(path, dpi=180); plt.close(fig); chart_paths.append(path)

    path = CHART_DIR / "carrier_performance.png"
    summary = data.groupby("carrier", as_index=False).agg(on_time_rate=("delivery_status", lambda s: (s == "On Time").mean()), average_cost=("shipping_cost", "mean"))
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ordered = summary.sort_values("on_time_rate")
    ax.barh(ordered["carrier"], ordered["on_time_rate"], color="#2A9D8F")
    ax.set(title="On-time delivery rate by carrier", xlabel="On-time rate", ylabel="Carrier")
    ax.set_xlim(0, 1); ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda value, _: f"{value:.0%}"))
    fig.tight_layout(); fig.savefig(path, dpi=180); plt.close(fig); chart_paths.append(path)

    path = CHART_DIR / "cost_distance_relationship.png"
    fig, ax = plt.subplots(figsize=(8, 4.5))
    sample = data.sample(500, random_state=42)
    for service, color in colors.items():
        subset = sample[sample["service_level"] == service]
        ax.scatter(subset["distance_km"], subset["shipping_cost"], alpha=0.5, s=18, color=color, label=service)
    coefficients = np.polyfit(data["distance_km"], data["shipping_cost"], 1)
    x_line = np.linspace(data["distance_km"].min(), data["distance_km"].max(), 100)
    ax.plot(x_line, np.polyval(coefficients, x_line), color="#111111", linewidth=2, label="Linear trend")
    ax.legend(title="Service level")
    ax.set(title="Shipping cost rises with distance", xlabel="Distance (km)", ylabel="Shipping cost")
    fig.tight_layout(); fig.savefig(path, dpi=180); plt.close(fig); chart_paths.append(path)

    path = CHART_DIR / "weekly_volume_delay.png"
    weekly = data.assign(week=data["order_date"].dt.to_period("W").dt.start_time).groupby("week").agg(volume=("carrier", "size"), average_delay=("delay_days", "mean")).reset_index()
    fig, ax1 = plt.subplots(figsize=(8, 4.5))
    ax1.plot(weekly["week"], weekly["volume"], marker="o", color="#1F4E79", label="Volume")
    ax1.set_ylabel("Shipment volume", color="#1F4E79"); ax1.set_xlabel("Week")
    ax2 = ax1.twinx(); ax2.plot(weekly["week"], weekly["average_delay"], color="#E76F51", linewidth=2, label="Average delay")
    ax2.set_ylabel("Average delay (days)", color="#E76F51")
    ax1.set_title("Weekly shipment volume and average delay")
    fig.tight_layout(); fig.savefig(path, dpi=180); plt.close(fig); chart_paths.append(path)

    path = CHART_DIR / "correlation_heatmap.png"
    numeric = data[["distance_km", "weight_kg", "delivery_time_days", "delay_days", "shipping_cost", "congestion_flag"]]
    fig, ax = plt.subplots(figsize=(8, 5))
    correlations = numeric.corr()
    image = ax.imshow(correlations, cmap="coolwarm", vmin=-1, vmax=1)
    ax.set_xticks(range(len(correlations.columns)), correlations.columns, rotation=45, ha="right")
    ax.set_yticks(range(len(correlations.columns)), correlations.columns)
    for row in range(len(correlations)):
        for column in range(len(correlations)):
            ax.text(column, row, f"{correlations.iloc[row, column]:.2f}", ha="center", va="center", fontsize=8)
    fig.colorbar(image, ax=ax, label="Correlation")
    ax.set_title("Correlation matrix of operational variables")
    fig.tight_layout(); fig.savefig(path, dpi=180); plt.close(fig); chart_paths.append(path)
    return chart_paths


def build_report(data, chart_paths):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75); section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(10.5); styles["Normal"].font.color.rgb = RGBColor(45, 55, 65)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Aptos Display"; styles["Title"].font.size = Pt(29); styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
    for name, size in (("Heading 1", 18), ("Heading 2", 13)):
        styles[name].font.name = "Aptos Display"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor(31, 78, 121)
    title = document.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.add_run("Advanced Data Analysis and\nVisualization in Logistics")
    sub = document.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = sub.add_run("Week 3 Analytical Findings Report"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(90, 105, 120)
    date = document.add_paragraph(); date.alignment = WD_ALIGN_PARAGRAPH.CENTER; date.add_run("Hypothetical shipment dataset | Python EDA and visual analytics | September 2026")
    document.add_paragraph()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph("This report analyses a reproducible hypothetical dataset of 1,000 logistics shipments over 180 days. The dataset represents carrier operations across four regions and three service levels, with distance, weight, congestion, delivery time, delay, cost, and delivery status variables. The analysis demonstrates how exploratory data analysis and visualisation convert raw operational records into evidence for service management, cost control, and capacity planning. Because the data is simulated, the numerical findings are illustrative rather than claims about a real carrier network.")
    document.add_paragraph("The results show the expected operational relationships built into the scenario: delivery time and shipping cost increase as distance, weight, and delay increase; economy service has a longer delivery-time profile; and congestion is associated with higher delay risk. Carrier comparison highlights a meaningful reliability gap, while weekly charts help distinguish demand volume from service deterioration. The findings support targeted actions such as reviewing the weakest carrier lanes, using service-level commitments that reflect actual transit distributions, and investigating congestion-sensitive routes before adding expensive capacity.")

    document.add_heading("1. Dataset Design and Methodology", level=1)
    document.add_paragraph("The hypothetical dataset models a shipment-level operational table. Each row is one shipment, allowing metrics to be aggregated by carrier, region, service level, and week. A fixed random seed makes the simulation repeatable, while controlled effects make the scenario analytically meaningful: longer distance and heavier weight increase cost and transit time; congestion and carrier-specific effects increase delay; and service level changes expected delivery time.")
    add_table(document, ["Variable", "Type", "Analytical purpose"], [
        ("order_date", "Date", "Weekly trend and seasonality-style grouping"),
        ("carrier, region, service_level", "Categorical", "Operational segment comparisons"),
        ("distance_km, weight_kg", "Numeric", "Network and load cost drivers"),
        ("congestion_flag", "Binary", "Potential delay driver"),
        ("delivery_time_days, delay_days", "Numeric", "Service performance outcomes"),
        ("shipping_cost", "Numeric", "Cost performance and relationship analysis"),
        ("delivery_status", "Categorical", "On-time rate calculation"),
    ], widths=[2.0, 1.4, 3.2])
    document.add_paragraph("The workflow includes descriptive statistics, grouped summaries, distribution analysis, bivariate relationships, a correlation matrix, and time aggregation. Visualizations were selected for their decision role: histograms show variation and service overlap, bars compare segments, scatterplots reveal cost drivers, a dual-axis trend chart shows volume against delay, and a heatmap provides a compact view of linear relationships.")

    document.add_heading("2. Exploratory Data Analysis", level=1)
    stats = data[["distance_km", "weight_kg", "delivery_time_days", "delay_days", "shipping_cost"]].describe().round(2)
    add_table(document, ["Measure", "Mean", "Median", "Std. dev.", "Minimum", "Maximum"], [
        (column.replace("_", " ").title(), f"{stats.loc['mean', column]:.2f}", f"{stats.loc['50%', column]:.2f}", f"{stats.loc['std', column]:.2f}", f"{stats.loc['min', column]:.2f}", f"{stats.loc['max', column]:.2f}") for column in stats.columns
    ], widths=[1.9, 1.0, 1.0, 1.0, 1.0, 1.0])
    overall_rate = (data["delivery_status"] == "On Time").mean()
    document.add_paragraph(f"Across the simulated records, the overall on-time delivery rate is {overall_rate:.1%}. Average delivery time is {data.delivery_time_days.mean():.2f} days, average delay is {data.delay_days.mean():.2f} days, and average shipping cost is {data.shipping_cost.mean():.2f} cost units. These values describe the scenario generated by the script and should be recalculated when real project data becomes available.")

    document.add_heading("3. Visualization Outputs and Interpretations", level=1)
    captions = [
        ("Delivery-time distribution", "The distribution separates service expectations: Economy shipments are shifted toward longer transit times, while Express is concentrated at shorter durations. This supports percentile-based promises rather than relying only on one average. Overlap still matters because variability can cause an Express shipment to perform like a slower service.", chart_paths[0]),
        ("Carrier performance comparison", "The carrier bar chart makes reliability differences easy to rank. A carrier with a lower on-time rate should be investigated by route, region, and service level before a network-wide conclusion is made. The most useful follow-up is a volume-weighted comparison with confidence intervals, not a decision based on a small segment alone.", chart_paths[1]),
        ("Cost and distance relationship", "The positive trend shows distance as a structural cost driver. Points above the fitted line may indicate heavier loads, premium service, congestion, or unusual charges. The chart helps analysts move from 'which carrier costs more?' to 'which cost components explain the difference?'.", chart_paths[2]),
        ("Weekly volume and delay", "The two series show whether service pressure rises with demand. A week with high volume and high average delay is a capacity or planning candidate; high delay at normal volume suggests a route, carrier, or disruption issue. Time aggregation is essential because a whole-period average can hide short operational failures.", chart_paths[3]),
        ("Correlation heatmap", "Delivery time, delay, and shipping cost show related movement in this scenario, while distance provides a direct structural link to cost and transit. Correlation is descriptive, not causal: it does not prove that changing one variable will produce a specific operational result. It is a screening tool for feature selection and investigation.", chart_paths[4]),
    ]
    for heading, interpretation, chart in captions:
        document.add_heading(heading, level=2)
        document.add_picture(str(chart), width=Inches(6.2))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(interpretation)

    document.add_heading("4. Analytical Insights and Recommendations", level=1)
    for item in [
        "Prioritise lane-level carrier review: segment the weakest carrier result by route and region, then compare volume, distance, service level, and delay severity before changing allocation.",
        "Use service distributions in customer promises: report median and 90th-percentile delivery time, not only the average, so operational variability is visible in service design.",
        "Treat distance and weight as cost levers: evaluate consolidation, zone-based pricing, and packaging or load planning where high cost is not explained by service level alone.",
        "Create congestion-sensitive exception monitoring: a high congestion flag combined with long distance or high weight can trigger proactive dispatch review.",
        "Separate demand pressure from execution failure: monitor weekly volume and delay together, then assign corrective ownership to capacity planning, carrier management, or network operations.",
    ]:
        add_bullet(document, item)
    document.add_paragraph("The recommendations are hypotheses generated from a controlled simulation. Real decisions require validation against actual shipment records, costs, contractual service targets, sample sizes, and operational constraints. Statistical significance, confidence intervals, and a review of confounding variables should be added before implementation.")

    document.add_heading("5. Python Analysis Script", level=1)
    document.add_paragraph("The accompanying script, docs/generate_week3_report.py, creates the data, calculates the summaries, exports five PNG visualizations, and embeds them into this report. The central analysis pattern is shown below.")
    add_code(document, '''data = build_dataset(seed=42, rows=1000)
weekly = (data.assign(week=data["order_date"].dt.to_period("W").dt.start_time)
             .groupby("week")
             .agg(volume=("carrier", "size"),
                  average_delay=("delay_days", "mean"))
             .reset_index())

carrier_summary = (data.groupby("carrier")
                     .agg(shipments=("delivery_status", "size"),
                          on_time_rate=("delivery_status", lambda s: (s == "On Time").mean()),
                          average_cost=("shipping_cost", "mean"),
                          average_delay=("delay_days", "mean"))
                     .sort_values("on_time_rate"))

correlations = data.select_dtypes("number").corr()
print(data.describe().round(2))
print(carrier_summary)
print(correlations.round(2))''')
    document.add_paragraph("For production use, the same visual functions should accept the cleaned project dataset rather than simulated data. Chart titles, axis units, filters, sample sizes, and date windows should be explicit so a dashboard user can interpret each result without guessing.")

    document.add_heading("6. Limitations and Conclusion", level=1)
    document.add_paragraph("The simulation is intentionally designed, so its relationships are not evidence about a real network. It does not include inventory position, promised delivery dates, weather, holidays, fuel prices, vehicle capacity, failed delivery attempts, or actual route geometry. Correlation does not establish causation, and a carrier average may conceal substantial lane-level variation. These limitations define the next analytical requirements rather than weakening the value of the exercise.")
    document.add_paragraph("The Week 3 analysis demonstrates how a logistics team can move from a structured shipment table to interpretable evidence. Descriptive statistics establish scale and variation; charts make segment differences and bottlenecks visible; correlations suggest relationships for deeper testing; and recommendations connect findings to operational actions. When actual data is loaded, this workflow can support KPI monitoring, carrier reviews, route diagnostics, and later predictive or optimisation work. The strongest decision-making outcome is not a single chart but a repeatable analytical process in which every visual has a clear question, a known limitation, and an accountable operational response.")

    document.add_heading("References", level=1)
    for reference in [
        "McKinney, W. Python for Data Analysis, 3rd edition. O'Reilly Media, 2022.",
        "Hunter, J. D. Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95, 2007.",
        "Waskom, M. Seaborn: statistical data visualization. Journal of Open Source Software, 6(60), 3021, 2021.",
        "Project sources: README.md, src/kpi.py, and notebooks/01_data_exploration.ipynb.",
    ]:
        add_bullet(document, reference)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("Logistics Data Science Project | Week 3")
    document.core_properties.title = "Advanced Data Analysis and Visualization in Logistics"
    document.core_properties.subject = "Week 3 logistics EDA and visual analytics report"
    document.core_properties.author = "Logistics Data Science Project"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    dataset = build_dataset()
    charts = save_charts(dataset)
    build_report(dataset, charts)