from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from generate_week1_report import add_bullet, add_code, add_table


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "week2_data_preprocessing_report.docx"


def build_report():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(45, 55, 65)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)
    styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
    for name, size in (("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(31, 78, 121)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Data Collection, Cleaning, and\nPreprocessing for Logistics Analysis")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Week 2 Data Preparation Report")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(90, 105, 120)
    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.add_run("A reproducible Python pipeline for trustworthy transport analytics | September 2026")
    document.add_paragraph()

    document.add_heading("Report Description", level=1)
    document.add_paragraph(
        "This report documents a practical Week 2 data preparation process for a logistics analytics project. It uses the publicly available New York City Taxi and Limousine Commission Trip Record Data as a reference dataset because it contains realistic transport-operation fields such as pickup and drop-off timestamps, locations, trip distance, passenger count, payment type, and fare amounts. These fields provide a useful public analogue for shipment and vehicle movement data: timestamps support duration calculations, locations support route analysis, and cost and distance fields support operational benchmarking. The report also explains how the same controls can be applied to the project repository's future shipment-level CSV.\n\n"
        "The proposed pipeline begins with source discovery and data-contract design. It then performs schema checks, type conversion, duplicate detection, missing-value profiling, categorical standardisation, date validation, outlier investigation, and numerical scaling. Each treatment is justified according to its effect on logistics decisions. For example, a missing delivery timestamp cannot be safely replaced with a mean because it changes the meaning of delivery-time KPIs, while a missing passenger count may be imputed only for a model if its absence is documented. Outliers are investigated against operational rules before removal, since a very long trip may represent a genuine disruption rather than an error.\n\n"
        "Python examples use pandas, NumPy, scikit-learn, and a reproducible audit log. They illustrate how to retain raw data, create a clean analytical table, distinguish invalid from unusual observations, and fit transformations on training data to avoid leakage. The report concludes that data quality is a decision-quality issue: inaccurate durations, costs, or route records can misdirect carrier evaluation, capacity planning, and predictive models. The current repository has not populated its raw logistics data, so this document presents an executable method and expected controls rather than fabricated numerical results."
    )

    document.add_heading("1. Data Collection Simulation", level=1)
    document.add_paragraph(
        "The reference source is the NYC TLC Trip Record Data portal. TLC publishes monthly trip records in Parquet format, together with a data dictionary and explanatory notes. A collection job would download a selected month, store the original file unchanged under data/raw/, record the source URL and retrieval date, and create a manifest containing file size, row count, checksum, and schema version. The data is used here as a transport-data reference, not as a claim that taxi trips are identical to parcel shipments."
    )
    add_table(document, ["Reference field", "Logistics interpretation", "Preparation need"], [
        ("tpep_pickup_datetime / tpep_dropoff_datetime", "Movement start and end events", "Parse timezone-aware datetimes; check order"),
        ("PULocationID / DOLocationID", "Origin and destination zones", "Treat as categorical route keys"),
        ("trip_distance", "Travel distance", "Validate positive values and plausible range"),
        ("fare_amount / total_amount", "Transport cost measures", "Validate currency fields; separate components"),
        ("passenger_count", "Load or occupancy proxy", "Check zero, missing, and extreme values"),
        ("payment_type", "Transaction category", "Standardise codes and preserve unknowns"),
    ], widths=[2.2, 2.4, 2.2])
    document.add_paragraph(
        "For the project shipment dataset, the equivalent collection contract should include Shipment ID, order and ship timestamps, delivery status, delivery duration, delay, shipping cost, carrier, route, origin, destination, weight, and service type where available. Each field needs a definition, unit, valid range, allowed categories, null policy, and business owner. Collection should be incremental and idempotent: rerunning a month must not silently duplicate records."
    )
    document.add_heading("Collection controls", level=2)
    for item in [
        "Version the source month and schema with the raw file; never overwrite the original download during cleaning.",
        "Record row counts, column names, file hash, retrieval date, and any provider warnings in a manifest.",
        "Keep personally identifying or address-level data out of the analytical table unless it is essential and access-controlled.",
        "Use a stable record identifier where possible; otherwise create a documented composite key for duplicate checks.",
    ]:
        add_bullet(document, item)

    document.add_heading("2. Data Quality Issues and Detection", level=1)
    document.add_paragraph(
        "Cleaning begins with measurement, not deletion. A profile is generated before any treatment so the team can quantify the starting condition and compare it with the cleaned output. The most important distinction is between invalid data, which violates a defined rule, and unusual data, which may be valid but deserves investigation. A record with a negative trip duration is invalid; a record with an unusually long but positive duration may be a genuine incident."
    )
    add_table(document, ["Issue", "Detection", "Preferred response"], [
        ("Missing values", "Null counts and missingness by segment", "Impute only when defensible; retain flags; exclude unusable targets"),
        ("Wrong types", "Schema and parse-error counts", "Convert explicitly; quarantine failed parses"),
        ("Duplicates", "Stable ID or composite-key counts", "Keep one record only when duplication is confirmed"),
        ("Invalid ranges", "Business-rule assertions", "Set invalid values to missing or quarantine for review"),
        ("Outliers", "IQR, robust z-score, and domain bounds", "Investigate, cap, transform, or retain with rationale"),
        ("Inconsistent categories", "Unique-value and frequency review", "Map known variants; preserve unknown category"),
    ], widths=[1.4, 2.6, 2.8])

    document.add_heading("3. Missing-Value Methodology", level=1)
    document.add_paragraph(
        "Missingness is analysed by field, time period, route, carrier, and outcome status. It can be missing completely at random, related to an observed operational condition, or caused by the outcome itself. The treatment must reflect the decision. A missing delivery date prevents a defensible delivery-duration calculation, so that record should not receive an invented duration; it can remain in an operational exception table. A missing numerical feature used by a predictive model may receive a training-set median, accompanied by a missingness indicator. A categorical field can use an explicit Unknown category rather than silently dropping rows."
    )
    add_code(document, '''import pandas as pd

raw = pd.read_parquet("data/raw/tlc_trip_records.parquet")
quality = pd.DataFrame({
    "dtype": raw.dtypes.astype(str),
    "missing_count": raw.isna().sum(),
    "missing_rate": raw.isna().mean(),
    "unique_values": raw.nunique(dropna=True),
}).sort_values("missing_rate", ascending=False)

# Preserve the signal that a value was absent before imputation.
raw["passenger_count_missing"] = raw["passenger_count"].isna()
raw["passenger_count"] = raw["passenger_count"].fillna(
    raw["passenger_count"].median()
)

# Delivery or trip duration is a target-derived measure: do not invent it.
raw["pickup_datetime"] = pd.to_datetime(raw["tpep_pickup_datetime"], errors="coerce")
raw["dropoff_datetime"] = pd.to_datetime(raw["tpep_dropoff_datetime"], errors="coerce")
raw["trip_duration_minutes"] = (
    raw["dropoff_datetime"] - raw["pickup_datetime"]
).dt.total_seconds() / 60
raw.loc[raw["trip_duration_minutes"].isna(), "duration_status"] = "unusable"
raw.loc[raw["trip_duration_minutes"].notna(), "duration_status"] = "usable"''')

    document.add_heading("4. Cleaning and Validation Pipeline", level=1)
    document.add_paragraph(
        "The pipeline produces three outputs: a raw immutable table, a cleaned analytical table, and a quality report. Every transformation is deterministic and recorded. Records removed from the analytical table are not destroyed; they are written to a quarantine file with a reason code such as duplicate_id, invalid_timestamp, negative_cost, or impossible_duration. This makes the process auditable and allows rules to be corrected without reacquiring the source."
    )
    add_code(document, '''import numpy as np

clean = raw.copy()
clean = clean.drop_duplicates()
clean["payment_type"] = clean["payment_type"].astype("string").str.strip().str.upper()
clean["trip_distance"] = pd.to_numeric(clean["trip_distance"], errors="coerce")
clean["total_amount"] = pd.to_numeric(clean["total_amount"], errors="coerce")

invalid = (
    clean["trip_duration_minutes"].notna()
    & (clean["trip_duration_minutes"] <= 0)
) | (clean["trip_distance"].notna() & (clean["trip_distance"] < 0))
quarantine = clean.loc[invalid].assign(reason="business_rule_failure")
clean = clean.loc[~invalid].copy()

assert clean["trip_duration_minutes"].dropna().gt(0).all()
assert clean["trip_distance"].dropna().ge(0).all()
clean.to_parquet("data/processed/clean_transport.parquet", index=False)
quarantine.to_parquet("data/processed/quarantine_transport.parquet", index=False)''')
    document.add_paragraph(
        "For the shipment project, validation rules should include ship date not after delivery date, non-negative shipping cost, delay consistent with the delivery status, and unique Shipment ID. Rules should be tested against a small sample first, then run across the full file with before-and-after row counts. A failure should stop the pipeline when a critical assumption is broken rather than silently producing partial output."
    )

    document.add_heading("5. Outlier Detection and Treatment", level=1)
    document.add_paragraph(
        "Outlier treatment is a modelling choice, not a synonym for deleting the largest values. Operational extremes can contain the exact disruptions that logistics leaders need to understand. The process therefore uses domain checks first, followed by statistical screening. IQR fences provide a transparent exploratory flag, while robust z-scores are less sensitive to extreme values than mean-based z-scores. A flagged record is retained unless it is demonstrably erroneous, and the chosen action is recorded."
    )
    add_code(document, '''def iqr_flags(series: pd.Series) -> pd.Series:
    q1, q3 = series.quantile([0.25, 0.75])
    spread = q3 - q1
    return (series < q1 - 1.5 * spread) | (series > q3 + 1.5 * spread)

clean["distance_outlier"] = iqr_flags(clean["trip_distance"].dropna())
clean["cost_log"] = np.log1p(clean["total_amount"].clip(lower=0))

# Cap only for a model-sensitive feature after reviewing flagged records.
upper = clean["cost_log"].quantile(0.99)
clean["cost_log_capped"] = clean["cost_log"].clip(upper=upper)''')
    document.add_paragraph(
        "The example illustrates a common implementation detail: when assigning a filtered Series back to a full DataFrame, indexes must align. In production, compute the flag on the same indexed non-null subset or initialise the full column carefully. More importantly, the report should state whether the outlier was removed, winsorised, transformed, or retained, and how that decision changes KPI denominators."
    )

    document.add_heading("6. Normalization and Feature Preparation", level=1)
    document.add_paragraph(
        "Scaling is not required for every analysis. Tree-based models and rule-based summaries are often insensitive to feature scale, while distance-based clustering, k-nearest neighbours, and regularised linear models can be strongly affected by it. StandardScaler centres and scales using mean and standard deviation; RobustScaler uses median and interquartile range and is preferable when legitimate extremes remain. Categorical variables should be encoded, and transformations must be fit on training data only."
    )
    add_code(document, '''from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, RobustScaler

numeric = ["trip_distance", "trip_duration_minutes", "total_amount"]
categorical = ["PULocationID", "DOLocationID", "payment_type"]
preprocessor = ColumnTransformer([
    ("numeric", Pipeline([
        ("impute", SimpleImputer(strategy="median", add_indicator=True)),
        ("scale", RobustScaler()),
    ]), numeric),
    ("categorical", Pipeline([
        ("impute", SimpleImputer(strategy="most_frequent")),
        ("encode", OneHotEncoder(handle_unknown="ignore")),
    ]), categorical),
])

# The fitted transformer belongs inside the model pipeline.
X_train_ready = preprocessor.fit_transform(train[numeric + categorical])
X_test_ready = preprocessor.transform(test[numeric + categorical])''')
    document.add_paragraph(
        "Fitting the preprocessor on the full dataset would allow test-set statistics to influence training and make later model performance look better than it is. Keeping the transformer in a pipeline also ensures the same treatment is applied when new shipments arrive."
    )

    document.add_heading("7. Reproducible Process Flow", level=1)
    add_table(document, ["Stage", "Action", "Evidence produced"], [
        ("Acquire", "Download source and write manifest", "Raw file, URL, retrieval date, hash"),
        ("Profile", "Inspect schema, nulls, categories, ranges", "Initial quality report"),
        ("Standardise", "Parse dates, names, units, and categories", "Typed staging table"),
        ("Validate", "Apply rules and quarantine failures", "Reason-coded exceptions"),
        ("Treat", "Impute, flag, transform, or retain with rationale", "Clean analytical table"),
        ("Scale", "Fit training-only transformations", "Model-ready feature matrix"),
        ("Verify", "Compare counts and distributions before/after", "Release checklist and sign-off"),
    ], widths=[1.2, 3.6, 2.4])
    document.add_paragraph(
        "The release checklist should include source-to-output row reconciliation, duplicate counts, missingness thresholds, valid date ranges, negative-value checks, category mapping counts, outlier review, and a sample of transformed records. The cleaned file is ready for Week 3 modelling only when an analyst and an operational data owner agree that exclusions and imputations preserve the intended business meaning."
    )

    document.add_heading("8. Reflection and Conclusion", level=1)
    document.add_paragraph(
        "Data quality directly controls the credibility of logistics decisions. A duplicated shipment can inflate volume and cost. A missing delivery timestamp can bias service-level calculations toward records that completed successfully. A unit error in distance or cost can make one carrier appear inefficient and can distort route clusters. Unchecked outliers can dominate averages, while careless deletion can hide disruptions that should drive resilience planning. These are not cosmetic defects; they alter what managers believe about capacity, performance, and risk."
    )
    document.add_paragraph(
        "The Week 2 outcome is a documented preprocessing contract that future analysts can rerun and audit. It preserves raw evidence, makes uncertainty visible, and ensures that modelling transformations are applied consistently. Once populated project shipment data is available, the same process will produce a clean dataset for the Week 1 KPIs and the planned delay-prediction and optimisation work. The recommended next step is a small pilot month, reviewed jointly by data and operations owners, before scaling collection across all carriers and periods."
    )

    document.add_heading("References", level=1)
    for reference in [
        "New York City Taxi and Limousine Commission. TLC Trip Record Data. https://www.nyc.gov/site/tlc/about/tlc-trip-record-data.page (accessed September 3, 2026).",
        "U.S. Bureau of Transportation Statistics. Freight Analysis Framework. https://faf.ornl.gov/faf5/ (accessed September 3, 2026).",
        "McKinney, W. Python for Data Analysis, 3rd edition. O'Reilly Media, 2022.",
        "Scikit-learn documentation. Imputation of missing values and preprocessing data. https://scikit-learn.org/stable/modules/impute.html and https://scikit-learn.org/stable/modules/preprocessing.html (accessed September 3, 2026).",
        "Project sources: README.md, src/kpi.py, src/data_cleaning.py, and notebooks/01_data_exploration.ipynb.",
    ]:
        add_bullet(document, reference)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Logistics Data Science Project | Week 2")
    document.core_properties.title = "Data Collection, Cleaning, and Preprocessing for Logistics Analysis"
    document.core_properties.subject = "Week 2 logistics data preparation report"
    document.core_properties.author = "Logistics Data Science Project"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()