from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "week1_strategic_planning_report.docx"


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=(255, 255, 255))
        shade(table.rows[0].cells[index], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def add_code(document, code):
    paragraph = document.add_paragraph()
    paragraph.style = "No Spacing"
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    run = paragraph.add_run(code.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(35, 45, 55)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EEF3F7")
    paragraph._p.get_or_add_pPr().append(shading)
    document.add_paragraph()


def add_bullet(document, text):
    document.add_paragraph(text, style="List Bullet")


def build_report():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(45, 55, 65)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)
    styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
    for name, size in (("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(31, 78, 121)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Strategic Planning and Data Exploration\nin Logistics")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Week 1 Strategic Planning Report")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(90, 105, 120)
    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.add_run("Applied shipment performance, cost, and delay analytics | September 2026")
    document.add_paragraph()

    document.add_heading("Submission Description", level=1)
    document.add_paragraph(
        "This report presents the strategic planning foundation for a logistics data science project focused on shipment reliability, cost control, and operational resource allocation. It defines a realistic parcel and freight-distribution scenario in which a logistics operator needs to understand why deliveries are late, which routes and carriers consume disproportionate cost, and where predictive and optimisation methods can create practical improvement. The report translates those business questions into measurable KPIs, including on-time delivery rate, average delivery time, average delay, average shipping cost, total shipping cost, and shipment volume. It also describes the data required to calculate each KPI and the controls needed before management decisions are made.\n\n"
        "The proposed workflow is deliberately reproducible in Python. It starts with source-data inventory and schema validation, continues through cleaning and exploratory analysis, and then progresses to delay prediction, shipment segmentation, and route or carrier decision support. Illustrative pandas, scikit-learn, and optimisation pseudocode shows how the plan can become an executable analysis. Public references from the World Bank Logistics Performance Index, the U.S. Bureau of Transportation Statistics Freight Analysis Framework, and scikit-learn documentation provide external context for the methodology.\n\n"
        "The repository currently contains the analysis structure, KPI module, and exploratory notebook, but no populated raw data and an empty processed CSV. Consequently, this Week 1 submission reports no invented numerical findings. Instead, it establishes a credible measurement contract and a roadmap for producing defensible findings when shipment records become available. The expected result is a decision-ready evidence base that helps operations leaders prioritise late-delivery interventions, capacity planning, carrier management, and future network optimisation."
    )

    document.add_heading("1. Background and Project Definition", level=1)
    document.add_paragraph(
        "The proposed organisation is a regional logistics operator serving business customers through a multi-carrier distribution network. Orders are collected from warehouses, assigned to carriers, transported through route and regional networks, and delivered to customers. Leadership currently sees the final delivery status and cost but lacks a consistent analytical view of the drivers behind performance differences. The Week 1 project creates that view from shipment-level records."
    )
    document.add_paragraph(
        "The central planning question is: how can reliable shipment data be converted into earlier warnings and better allocation of transport resources without sacrificing service quality? The initial unit of analysis is one shipment. Later phases can aggregate to route, carrier, warehouse, product category, destination region, and week. This structure supports both operational questions, such as which shipments need intervention today, and strategic questions, such as which carrier or lane should receive more volume."
    )
    document.add_heading("Business questions", level=2)
    for item in [
        "Which routes, carriers, regions, and time periods have the weakest delivery reliability?",
        "How do shipping cost and delivery time vary across operational segments?",
        "Which observable shipment features provide an early signal of delay?",
        "How should limited vehicles, carrier capacity, or expedited-service budget be allocated?",
    ]:
        add_bullet(document, item)

    document.add_heading("2. Objectives and Key Performance Indicators", level=1)
    document.add_paragraph(
        "The project objectives are to establish a trustworthy data foundation, quantify current service and cost performance, identify operational drivers of delay, and produce recommendations that can be monitored over time. The KPI definitions below follow the reusable functions already implemented in src/kpi.py."
    )
    add_table(document, ["KPI", "Definition", "Decision use"], [
        ("On-time delivery rate", "On-time delivered shipments / delivered shipments", "Service level, carrier review, escalation threshold"),
        ("Average delivery time", "Mean Delivery Time (days)", "Promise design, customer expectation, lane comparison"),
        ("Average delay", "Mean positive Delay (days) among delayed records", "Severity of exceptions and recovery planning"),
        ("Average shipping cost", "Mean Shipping Cost per shipment", "Cost benchmarking and mode/carrier mix"),
        ("Total shipping cost", "Sum of Shipping Cost", "Budget impact and savings opportunity"),
        ("Shipment volume", "Number of shipment records", "Demand scale and denominator for all rates"),
    ], widths=[1.5, 2.7, 2.6])
    document.add_paragraph(
        "KPI governance matters as much as calculation. Every dashboard should show the measurement period, shipment population, missing-value treatment, and denominator. On-time rate should be reported only for delivered shipments when that status distinction is available. Results should also be sliced by route and carrier so an acceptable overall average does not hide a high-risk lane."
    )

    document.add_heading("3. Literature and Data Research", level=1)
    document.add_paragraph(
        "External logistics research supports a layered measurement strategy. The World Bank Logistics Performance Index provides a cross-country framework built around customs, infrastructure, international shipments, logistics competence, timeliness, and tracking and tracing. Although this project operates at shipment level, those dimensions reinforce the importance of treating reliability, visibility, and operational capability as connected rather than isolated outcomes."
    )
    document.add_paragraph(
        "The U.S. Bureau of Transportation Statistics Freight Analysis Framework combines freight origin-destination flows with commodity and mode information. It is useful context for later network analysis and scenario planning, especially when internal shipment data needs an external benchmark or a broader view of demand by geography and mode."
    )
    document.add_paragraph(
        "The modelling plan uses established data science techniques. Regression or classification estimates delay risk from features such as carrier, route, distance, weight, service type, weekday, and warehouse. Clustering groups similar lanes or shipment profiles so interventions can be tailored. Optimisation converts forecasts and business constraints into recommended assignments, such as minimising cost while meeting a service-level target. These methods should support human decisions, with clear validation and an audit trail, rather than replace operational judgement without review."
    )
    document.add_heading("Proposed data inventory", level=2)
    add_table(document, ["Data source", "Likely fields", "Purpose"], [
        ("Shipment transactions", "Shipment ID, order date, ship date, delivery date, status", "Core KPI population and timing"),
        ("Transport and carrier records", "Carrier, service type, mode, cost, tracking events", "Carrier and cost performance"),
        ("Network reference", "Origin, destination, route, distance, region", "Lane comparison and optimisation"),
        ("Warehouse and order data", "Product, weight, volume, warehouse, priority", "Capacity and handling drivers"),
        ("External context", "Weather, holidays, traffic, public freight flows", "Explanatory features and benchmarks"),
    ], widths=[1.6, 2.8, 2.4])
    document.add_paragraph(
        "The current repository expects a processed file at data/processed/clean_logistics.csv with fields including Shipping Cost, Delivery Status, Delivery Time (days), and Delay (days). Raw files are not yet present, so field names, units, and status categories must be confirmed during data onboarding rather than assumed."
    )

    document.add_heading("4. End-to-End Strategic Roadmap", level=1)
    add_table(document, ["Phase", "Activities", "Output and gate"], [
        ("1. Collect and profile", "Acquire raw exports, document owners, inspect rows, types, ranges, and identifiers.", "Data dictionary; access and completeness check"),
        ("2. Clean and validate", "Parse dates, standardise categories, remove duplicates, handle missing values, test business rules.", "Analysis-ready table; validation log"),
        ("3. Explore and segment", "Summarise KPIs by carrier, route, region, product, and time; inspect distributions and outliers.", "EDA notebook; priority segments"),
        ("4. Model delay risk", "Create a time-aware train/test split; compare baseline and interpretable models; calibrate probabilities.", "Validated risk score and driver analysis"),
        ("5. Optimise decisions", "Use costs, capacity, risk, and service constraints to evaluate carrier or route scenarios.", "Scenario recommendations; sensitivity analysis"),
        ("6. Deploy and monitor", "Publish KPI definitions, refresh pipeline, alerts, ownership, and drift checks.", "Dashboard-ready dataset and operating cadence"),
    ], widths=[1.3, 3.8, 1.7])
    document.add_paragraph(
        "The immediate Week 1 gate is not a model score. It is evidence that the data contract and KPI logic are explicit, reproducible, and safe to extend. Week 2 should begin only after source columns and status semantics are confirmed. Model evaluation should use a later time period as the test set to avoid leaking future information into training."
    )

    document.add_heading("5. Python Code Illustrations", level=1)
    document.add_heading("Load, validate, and calculate KPIs", level=2)
    add_code(document, '''from pathlib import Path
import sys
import pandas as pd

ROOT = Path.cwd()
data = pd.read_csv(ROOT / "data/processed/clean_logistics.csv")
required = {"Shipping Cost", "Delivery Status", "Delivery Time (days)", "Delay (days)"}
missing = required - set(data.columns)
if data.empty or missing:
    raise ValueError(f"Data is not ready: empty={data.empty}, missing={sorted(missing)}")

sys.path.insert(0, str(ROOT / "src"))
from kpi import calculate_kpis

kpis = pd.Series(calculate_kpis(data), name="value")
print(kpis)''')
    document.add_heading("Exploratory segmentation", level=2)
    add_code(document, '''# Compare reliability and cost only where the fields are available.
segment = (data.groupby("Carrier", dropna=False)
             .agg(shipments=("Delivery Status", "size"),
                  avg_cost=("Shipping Cost", "mean"),
                  avg_delay=("Delay (days)", "mean"))
             .sort_values("avg_delay", ascending=False))

status_rate = data.assign(on_time=data["Delivery Status"].eq("On Time"))
status_rate = status_rate.groupby("Carrier")["on_time"].mean()
segment["on_time_rate"] = status_rate
print(segment.sort_values("on_time_rate"))''')
    document.add_heading("Delay prediction and clustering", level=2)
    add_code(document, '''from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.cluster import KMeans

features = ["Carrier", "Route", "Shipping Cost", "Delivery Time (days)"]
X = data[features]
y = data["Delivery Status"].eq("Delayed").astype(int)
categorical = ["Carrier", "Route"]
numeric = ["Shipping Cost", "Delivery Time (days)"]
preprocess = ColumnTransformer([
    ("categorical", OneHotEncoder(handle_unknown="ignore"), categorical),
    ("numeric", SimpleImputer(strategy="median"), numeric),
])
model = Pipeline([("prep", preprocess), ("classifier", LogisticRegression(max_iter=1000))])
model.fit(X, y)  # In production, fit only on an earlier time window.
data["delay_risk"] = model.predict_proba(X)[:, 1]

scaled_cost = StandardScaler().fit_transform(data[["Shipping Cost", "Delivery Time (days)"]].fillna(0))
data["shipment_cluster"] = KMeans(n_clusters=3, random_state=42, n_init=10).fit_predict(scaled_cost)''')
    document.add_paragraph(
        "The examples are intentionally small illustrations, not production-ready model training. Before deployment, the team should add temporal validation, class-imbalance checks, baseline comparisons, precision-recall reporting, probability calibration, fairness review by operational segment, and a feature-availability check that prevents post-delivery fields from leaking into an early warning model."
    )
    document.add_heading("Illustrative constrained assignment objective", level=2)
    document.add_paragraph(
        "For a later optimisation phase, let x(i,c) equal 1 when shipment or lane i is assigned to carrier c. A simple objective is to minimise expected cost plus a penalty for predicted delay, subject to carrier capacity and a minimum service level. The final coefficients must be agreed with operations and tested through scenarios before any automated decision is made."
    )
    add_code(document, '''# Pseudocode for a linear assignment model
minimise sum(cost[i, c] * x[i, c] + risk_penalty * delay_risk[i, c] * x[i, c]
            for i in shipments for c in carriers)
subject to sum(x[i, c] for c in carriers) == 1       for each shipment i
           sum(volume[i] * x[i, c] for i in shipments) <= capacity[c]
           expected_on_time(x) >= service_target
           x[i, c] in {0, 1}''')

    document.add_heading("6. Risks, Ethics, and Quality Controls", level=1)
    for item in [
        "Data quality risk: incomplete timestamps or inconsistent status labels can distort both rates and model targets. Keep a validation log and publish record counts after every transformation.",
        "Leakage risk: delivery completion fields can make a model appear accurate while being unavailable at dispatch time. Define the prediction moment before feature engineering.",
        "Operational risk: a low-cost assignment can increase customer harm if service constraints are weak. Use explicit service targets, exception handling, and human approval for high-impact changes.",
        "Privacy and governance risk: restrict access to customer and address data, minimise personally identifiable information, and retain only fields needed for the stated decision.",
        "Drift risk: carrier mix, seasons, fuel prices, and network conditions change. Monitor KPI definitions, input distributions, calibration, and realised outcomes after deployment.",
    ]:
        add_bullet(document, item)

    document.add_heading("7. Expected Outcomes and Conclusion", level=1)
    document.add_paragraph(
        "By the end of the planned four-week cycle, the project should provide a trusted shipment-level dataset, repeatable KPI calculations, segment-level performance evidence, an interpretable delay-risk baseline, and scenario-based recommendations for carrier and route decisions. The immediate value is visibility: leaders can identify where service failures and cost concentration occur instead of relying on a single overall average. The next value is prioritisation: risk scores and clusters can direct attention to shipments and lanes where intervention is most likely to matter. The longer-term value is resource allocation: optimisation can compare feasible alternatives while preserving service commitments."
    )
    document.add_paragraph(
        "This Week 1 report therefore treats data exploration as strategic planning, not merely chart production. It establishes what will be measured, why it matters, how it will be validated, and how analysis will connect to decisions. Because the repository does not yet contain source records, numerical findings are intentionally deferred. That limitation is itself a useful result: populating and validating the source data is the next controlled milestone before performance claims are made."
    )

    document.add_heading("References", level=1)
    references = [
        "World Bank. Logistics Performance Index. https://lpi.worldbank.org/ (accessed September 3, 2026).",
        "U.S. Bureau of Transportation Statistics. Freight Analysis Framework. https://faf.ornl.gov/faf5/ (accessed September 3, 2026).",
        "Pedregosa, F. et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830, 2011. https://jmlr.org/papers/v12/pedregosa11a.html.",
        "Scikit-learn documentation. Model selection and evaluation. https://scikit-learn.org/stable/model_selection.html (accessed September 3, 2026).",
        "Project sources: README.md, src/kpi.py, src/data_cleaning.py, and notebooks/01_data_exploration.ipynb.",
    ]
    for reference in references:
        add_bullet(document, reference)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Logistics Data Science Project | Week 1")
    document.core_properties.title = "Strategic Planning and Data Exploration in Logistics"
    document.core_properties.subject = "Week 1 logistics data science strategic planning report"
    document.core_properties.author = "Logistics Data Science Project"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()