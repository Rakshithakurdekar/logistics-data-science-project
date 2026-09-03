from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from generate_week1_report import add_bullet, add_code, add_table


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "week4_predictive_modeling_optimization_report.docx"


def simulate_data(seed=42, rows=1200):
    rng = np.random.default_rng(seed)
    dates = pd.date_range("2025-01-01", periods=240, freq="D")
    data = pd.DataFrame({
        "order_date": rng.choice(dates, rows),
        "distance_km": np.clip(rng.gamma(3, 55, rows), 10, 500),
        "weight_kg": np.clip(rng.lognormal(2.6, 0.6, rows), 1, 65),
        "congestion_index": rng.uniform(0, 1, rows),
        "priority": rng.choice([0, 1], rows, p=[0.72, 0.28]),
        "carrier": rng.choice(["NorthStar", "SwiftHaul", "MetroFreight", "BlueLine"], rows),
    })
    carrier_effect = data["carrier"].map({"NorthStar": 0.15, "SwiftHaul": 0.45, "MetroFreight": 0.3, "BlueLine": 0.75})
    data["delivery_time_days"] = np.maximum(0.5, 1.1 + data.distance_km / 185 + data.weight_kg / 60 + data.congestion_index * 1.2 - data.priority * 0.35 + carrier_effect + rng.normal(0, 0.38, rows))
    data["delivery_time_days"] = data["delivery_time_days"].round(2)
    return data.sort_values("order_date").reset_index(drop=True)


def one_hot(data, columns):
    encoded = pd.get_dummies(data[columns], columns=["carrier"], dtype=float)
    return encoded.astype(float)


def fit_and_evaluate(data):
    split = int(len(data) * 0.8)
    train, test = data.iloc[:split], data.iloc[split:]
    features = ["distance_km", "weight_kg", "congestion_index", "priority", "carrier"]
    X_train = one_hot(train, features)
    X_test = one_hot(test, features).reindex(columns=X_train.columns, fill_value=0)
    X_train = np.column_stack([np.ones(len(X_train)), X_train.to_numpy()])
    X_test = np.column_stack([np.ones(len(X_test)), X_test.to_numpy()])
    coefficients = np.linalg.lstsq(X_train, train["delivery_time_days"].to_numpy(), rcond=None)[0]
    predictions = X_test @ coefficients
    actual = test["delivery_time_days"].to_numpy()
    mae = float(np.mean(np.abs(actual - predictions)))
    rmse = float(np.sqrt(np.mean((actual - predictions) ** 2)))
    r2 = float(1 - np.sum((actual - predictions) ** 2) / np.sum((actual - actual.mean()) ** 2))
    baseline = np.full(len(test), train["delivery_time_days"].mean())
    baseline_mae = float(np.mean(np.abs(actual - baseline)))
    return train, test, predictions, {"mae": mae, "rmse": rmse, "r2": r2, "baseline_mae": baseline_mae}, coefficients


def allocation_scenario(data):
    recent = data.tail(80).copy()
    capacity = {"NorthStar": 28, "SwiftHaul": 18, "MetroFreight": 20, "BlueLine": 14}
    carrier_cost = {"NorthStar": 42, "SwiftHaul": 38, "MetroFreight": 40, "BlueLine": 35}
    carrier_risk = {"NorthStar": 0.08, "SwiftHaul": 0.16, "MetroFreight": 0.12, "BlueLine": 0.22}
    recent["risk"] = (0.25 + recent["congestion_index"] * 0.45 + recent["distance_km"] / 900 + recent["weight_kg"] / 180).clip(0, 0.95)
    assignments = []
    remaining = capacity.copy()
    for _, shipment in recent.sort_values("risk", ascending=False).iterrows():
        choices = [carrier for carrier, slots in remaining.items() if slots > 0]
        selected = min(choices, key=lambda carrier: carrier_cost[carrier] + 85 * (shipment["risk"] + carrier_risk[carrier]))
        remaining[selected] -= 1
        assignments.append(selected)
    recent["recommended_carrier"] = assignments
    cost = sum(carrier_cost[c] for c in assignments)
    expected_risk = np.mean([recent.iloc[i]["risk"] + carrier_risk[c] for i, c in enumerate(assignments)])
    return recent, cost, expected_risk


def build_report(data, metrics, allocation):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65); section.left_margin = Inches(0.75); section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(10.5); styles["Normal"].font.color.rgb = RGBColor(45, 55, 65); styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Aptos Display"; styles["Title"].font.size = Pt(29); styles["Title"].font.color.rgb = RGBColor(31, 78, 121)
    for name, size in (("Heading 1", 18), ("Heading 2", 13)):
        styles[name].font.name = "Aptos Display"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor(31, 78, 121)
    title = document.add_paragraph(style="Title"); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.add_run("Predictive Modeling and\nOptimization in Logistics Systems")
    sub = document.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = sub.add_run("Week 4 Predictive Analytics Report"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(90, 105, 120)
    date = document.add_paragraph(); date.alignment = WD_ALIGN_PARAGRAPH.CENTER; date.add_run("Delivery-time forecasting and capacity allocation | September 2026")
    document.add_paragraph()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(f"This report develops a predictive and optimization approach for logistics operations using a deterministic hypothetical dataset of {len(data):,} shipments. The prediction target is delivery time in days, selected because it directly affects customer promises, service-level monitoring, and exception planning. Features include distance, weight, congestion, priority, carrier, and order date. A time-ordered 80/20 split evaluates whether a model trained on earlier shipments generalizes to later shipments, reducing the risk of future information leakage.")
    document.add_paragraph(f"An interpretable linear regression is compared with a historical-mean baseline. On the held-out period, the regression achieves MAE {metrics['mae']:.2f} days, RMSE {metrics['rmse']:.2f} days, and R-squared {metrics['r2']:.2f}; the baseline MAE is {metrics['baseline_mae']:.2f} days. The model is therefore useful as a planning baseline in this simulation, while its residual error means predictions should support decisions rather than replace operational review. A capacity allocation exercise then assigns 80 upcoming shipments across four carriers, respecting capacity limits and balancing carrier cost against predicted risk.")

    document.add_heading("1. Problem Definition and Data Simulation", level=1)
    document.add_paragraph("A regional logistics operator needs an earlier estimate of delivery time before dispatch. Better forecasts can improve customer promise dates, identify shipments needing proactive intervention, and help allocate scarce carrier capacity. The target is continuous delivery_time_days, while the operational decision is which carrier should receive each shipment under capacity and service constraints.")
    add_table(document, ["Feature", "Type", "Role in the model"], [
        ("distance_km", "Numeric", "Expected transit and cost driver"),
        ("weight_kg", "Numeric", "Load and handling complexity"),
        ("congestion_index", "Numeric 0-1", "Road/network condition signal"),
        ("priority", "Binary", "Service urgency indicator"),
        ("carrier", "Categorical", "Carrier-specific operational effect"),
        ("order_date", "Date", "Time ordering for validation and future features"),
        ("delivery_time_days", "Continuous target", "Forecasted logistics metric"),
    ], widths=[1.8, 1.5, 3.3])
    document.add_paragraph("The data is simulated with a fixed seed so the experiment can be reproduced. The generator embeds realistic relationships: longer and heavier shipments take longer, congestion increases transit time, priority reduces expected time, and carrier effects differ. These relationships are illustrative and must not be interpreted as measured performance of real companies.")

    document.add_heading("2. Model Selection and Implementation", level=1)
    document.add_paragraph("The primary model is multiple linear regression with one-hot encoded carrier categories. It is selected as an interpretable first model: coefficients can be inspected by analysts, training is fast, and its additive structure creates a useful benchmark. A historical-mean predictor is included as a naive baseline. A tree ensemble such as RandomForestRegressor would be a sensible next candidate if nonlinear relationships or interactions materially reduce error, but it should earn that complexity through time-aware validation and operational interpretability.")
    add_code(document, '''# Feature preparation and time-aware holdout
data = data.sort_values("order_date")
split = int(len(data) * 0.8)
train, test = data.iloc[:split], data.iloc[split:]
X_train = pd.get_dummies(train[features], columns=["carrier"], dtype=float)
X_test = pd.get_dummies(test[features], columns=["carrier"], dtype=float)
X_test = X_test.reindex(columns=X_train.columns, fill_value=0)

# Equivalent scikit-learn implementation for production
from sklearn.compose import ColumnTransformer
from sklearn.linear_model import Ridge
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler

model = Pipeline([
    ("features", ColumnTransformer([
        ("numeric", StandardScaler(), numeric_features),
        ("categorical", OneHotEncoder(handle_unknown="ignore"), categorical_features),
    ])),
    ("regressor", Ridge(alpha=1.0)),
])
model.fit(train[features], train[target])
predictions = model.predict(test[features])''')

    document.add_heading("3. Evaluation and Validation", level=1)
    add_table(document, ["Metric", "Result", "Interpretation"], [
        ("MAE", f"{metrics['mae']:.2f} days", "Average absolute forecast error; easy to communicate"),
        ("RMSE", f"{metrics['rmse']:.2f} days", "Penalises larger misses more strongly"),
        ("R-squared", f"{metrics['r2']:.2f}", "Share of held-out target variation explained"),
        ("Baseline MAE", f"{metrics['baseline_mae']:.2f} days", "Reference error from predicting the train mean"),
    ], widths=[1.5, 1.4, 3.7])
    document.add_paragraph("MAE is the primary operational metric because a mean error in days maps naturally to promise and planning decisions. RMSE highlights whether a small number of large misses remain, and R-squared provides context about explained variation but should not be used alone. A model is valuable only if it beats the baseline consistently and its error is acceptable for the decision being supported.")
    document.add_heading("Validation plan", level=2)
    for item in [
        "Use expanding-window or rolling time-series cross-validation once more historical periods are available.",
        "Tune Ridge regularization or tree depth only inside the training period; never use the final test period for repeated model selection.",
        "Report metrics overall and by carrier, region, distance band, and service level so weak segments are not hidden by aggregation.",
        "Check residuals for systematic underprediction on long, heavy, congested, or priority shipments.",
        "Monitor prediction drift and calibration after deployment; compare forecast error with realised delivery time each refresh cycle.",
    ]:
        add_bullet(document, item)

    document.add_heading("4. Optimization Strategy", level=1)
    document.add_paragraph("Prediction estimates the operational consequences; optimization chooses an action under constraints. The simulated allocation assigns 80 upcoming shipments to four carriers. Each carrier has a capacity limit and an indicative cost. The assignment objective balances direct cost with a penalty for expected delay risk. In production, the risk term would come from a validated model and the constraints would include contractual service levels, vehicle compatibility, geography, cutoff times, and fairness across carrier commitments.")
    recent, cost, expected_risk = allocation
    counts = recent["recommended_carrier"].value_counts()
    add_table(document, ["Carrier", "Capacity", "Recommended volume", "Indicative cost"], [
        (carrier, capacity, int(counts.get(carrier, 0)), f"{carrier_cost:.0f} per shipment") for carrier, capacity, carrier_cost in [("NorthStar", 28, 42), ("SwiftHaul", 18, 38), ("MetroFreight", 20, 40), ("BlueLine", 14, 35)]
    ], widths=[1.7, 1.3, 1.8, 1.8])
    document.add_paragraph(f"In the simulated recommendation, the 80 shipments are assigned within the stated capacities. The indicative allocation cost is {cost:.0f} cost units and the combined expected risk score averages {expected_risk:.2f}. These are scenario outputs, not real savings or service guarantees. A baseline comparison should be added before implementation, such as current carrier shares or lowest-cost-only assignment, followed by sensitivity tests for capacity, penalty weights, and forecast error.")
    add_code(document, '''# Binary assignment formulation
minimise sum((cost[i, c] + risk_penalty * predicted_delay[i, c]) * x[i, c]
             for i in shipments for c in carriers)
subject to sum(x[i, c] for c in carriers) == 1       for every shipment i
           sum(volume[i] * x[i, c] for i in shipments) <= capacity[c]
           weighted_on_time(x) >= service_target
           x[i, c] in {0, 1}

# In practice, solve with scipy.optimize.milp, OR-Tools, or PuLP.
# Compare the optimized plan with current allocation and stress scenarios.''')

    document.add_heading("5. Recommendations", level=1)
    for item in [
        "Adopt the regression as an explainable planning baseline and establish an error threshold in days with operations and customer-service owners.",
        "Use predicted delivery time to flag shipments likely to miss the promise window early enough for carrier or customer intervention.",
        "Allocate capacity with an explicit service penalty, rather than minimising transport cost alone; publish the trade-off so managers can challenge assumptions.",
        "Review model performance by carrier and lane. A good overall MAE can coexist with unacceptable errors in one region or service class.",
        "Run what-if scenarios before changing contracts: vary capacity, congestion, risk penalty, and service target, then assess cost and reliability together.",
    ]:
        add_bullet(document, item)

    document.add_heading("6. Limitations and Conclusion", level=1)
    document.add_paragraph("The dataset is hypothetical and generated from known relationships, so the reported metrics demonstrate workflow execution rather than real-world model quality. The experiment omits inventory availability, weather, holidays, traffic incidents, route geometry, promised dates, failed attempts, and actual carrier contract terms. The optimization example uses a transparent greedy assignment to illustrate the objective; a production system should use a validated mixed-integer solver and compare multiple feasible scenarios.")
    document.add_paragraph("This Week 4 exercise shows how predictive modeling and optimization work together in a logistics system. A time-aware forecast provides an estimate of delivery-time risk, evaluation metrics quantify its usefulness, and constrained optimization turns that signal into an allocation plan. The recommended path is staged: validate the data and model on real historical shipments, establish operational and ethical guardrails, compare optimized plans with current practice, pilot on a limited lane, and monitor realised service and cost outcomes before scaling.")

    document.add_heading("References", level=1)
    for reference in [
        "Hyndman, R. J. and Athanasopoulos, G. Forecasting: Principles and Practice, 3rd edition, OTexts, 2021.",
        "Scikit-learn documentation. Linear models, model evaluation, and time-related cross-validation. https://scikit-learn.org/stable/ (accessed September 3, 2026).",
        "Google OR-Tools documentation. Vehicle routing and linear optimization. https://developers.google.com/optimization (accessed September 3, 2026).",
        "Project sources: README.md, src/kpi.py, and the Week 1-3 project reports.",
    ]:
        add_bullet(document, reference)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("Logistics Data Science Project | Week 4")
    document.core_properties.title = "Predictive Modeling and Optimization in Logistics Systems"
    document.core_properties.subject = "Week 4 logistics forecasting and optimization report"
    document.core_properties.author = "Logistics Data Science Project"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    dataset = simulate_data()
    _, _, _, evaluation, _ = fit_and_evaluate(dataset)
    allocation_result = allocation_scenario(dataset)
    build_report(dataset, evaluation, allocation_result)